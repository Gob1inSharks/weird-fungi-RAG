import pandas as pd
from datetime import datetime

import docx

from docx import Document
from fpdf import FPDF

df=pd.read_excel("assets/feed.xlsx"
                ,'sheet1'
                 )
teacher="谢凯年"

dfx = df[(df['上课老师'] == teacher) & (df['上课状态'] == '已完成')]

def getStudentInfo(startDate,endDate):

    studentData = []
    studentData.append('# 学生上课内容\n')

    for index, row in dfx.iterrows():

        currentDate = datetime.strptime(row[9],'%Y-%m-%d %H:%M:%S')

        if startDate < currentDate < endDate:

            name = row[1]
            date = currentDate.strftime("%B %d, %Y")
            content = ''.join(i for i in row[10].replace('\n','').replace('.',' 。').replace('特点：',' 。') if not i.isdigit())
            
            info = '## **{0}**的**{1}**课程的内容\n{2}\n'.format(name,date,content)

            studentData.append(info)

    return studentData

def writeToWord(data):

    doc = docx.Document()
    doc.add_paragraph('Class Info Q&A')
    for paragraph in data:
        doc.add_paragraph(paragraph)

    doc.save("assets/feed.docx")

    return doc

def writeToMarkdown(data):

    with open('assets/feed.md','w') as mdFile:

        for paragraph in data:
            
            mdFile.write(paragraph)
            mdFile.write('\n')

def export2PDF():

    doc = Document("assets/feed.docx")
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('SimSun','','assets/SIMSUN.ttf',uni=True)
    for paragraph in doc.paragraphs:
        pdf.set_font("SimSun", size=12)
        pdf.cell(200, 10, txt=paragraph.text, ln=True)
    pdf.output("assets/feed.pdf")

def run(startTime,endTime):
    writeToMarkdown(getStudentInfo(startTime,endTime))

if __name__ == '__main__':

    startTime = datetime(2024,1,1,0,0,0)
    endTime = datetime(2024,5,2,0,0,0)

    writeToMarkdown(getStudentInfo(startTime,endTime))
